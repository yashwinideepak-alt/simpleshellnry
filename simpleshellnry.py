import streamlit as st
import subprocess
import shlex
import os
import threading
from docx import Document
import base64

st.set_page_config(page_title="Simple Shell", page_icon="💻", layout="centered")

st.title("💻 MySimpleShell - Interactive Web Shell")
st.write("Run shell commands with process execution, piping, redirection, and background support.")


# ---------- Command Handling ----------
def run_pipeline(cmds):
    """Handles piped commands"""
    processes = []
    prev_process = None

    for cmd in cmds:
        args = shlex.split(cmd)
        if prev_process is None:
            process = subprocess.Popen(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        else:
            process = subprocess.Popen(args, stdin=prev_process.stdout, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        prev_process = process
        processes.append(process)

    output, error = processes[-1].communicate()
    return output.decode() if output else error.decode()


def handle_command(cmd, background=False):
    """Detects operation type and executes accordingly"""
    try:
        # Piping
        if "|" in cmd:
            cmds = [c.strip() for c in cmd.split("|")]
            output = run_pipeline(cmds)
            return f"[Piping]\n{output}"

        # Append Redirection >>
        elif ">>" in cmd:
            parts = cmd.split(">>")
            main_cmd, filename = parts[0].strip(), parts[1].strip()
            args = shlex.split(main_cmd)
            with open(filename, "a") as f:
                subprocess.run(args, stdout=f, stderr=subprocess.PIPE)
            return f"[Output Redirection (Append)] Output appended to {filename}"

        # Output Redirection >
        elif ">" in cmd:
            parts = cmd.split(">")
            main_cmd, filename = parts[0].strip(), parts[1].strip()
            args = shlex.split(main_cmd)
            with open(filename, "w") as f:
                subprocess.run(args, stdout=f, stderr=subprocess.PIPE)
            return f"[Output Redirection] Output written to {filename}"

        # Input Redirection <
        elif "<" in cmd:
            parts = cmd.split("<")
            main_cmd, filename = parts[0].strip(), parts[1].strip()
            args = shlex.split(main_cmd)
            with open(filename, "r") as f:
                result = subprocess.run(args, stdin=f, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            return f"[Input Redirection]\n{result.stdout.decode()}"

        # Background Process
        elif background:
            def background_run():
                subprocess.Popen(shlex.split(cmd))
            threading.Thread(target=background_run).start()
            return f"[Process Creation (Background)] {cmd}"

        # Normal Command
        else:
            result = subprocess.run(shlex.split(cmd), stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            return f"[Process Execution]\n{result.stdout.decode() or result.stderr.decode()}"

    except Exception as e:
        return f"Error: {str(e)}"


# ---------- Streamlit UI ----------
st.subheader("🧠 Execute Shell Command")

command = st.text_input("Enter command:", placeholder="e.g., ls | grep py")
background = st.checkbox("Run in background (&)")
run_button = st.button("Run")

if run_button:
    if command.strip():
        output = handle_command(command.strip(), background)
        st.code(output or "(no output)", language="bash")
    else:
        st.warning("Please enter a command.")


# ---------- File Creation with Word Export ----------
st.subheader("📄 Create a New File")

filename = st.text_input("Filename:", "notes.txt")
content = st.text_area("File content:", "Hello from my simple shell!")

if st.button("Create File"):
    try:
        # Save as .txt file
        with open(filename, "w") as f:
            f.write(content)
        st.success(f"✅ File '{filename}' created successfully!")

        # Create Word (.docx) version
        doc = Document()
        doc.add_heading(filename.replace('.txt', ''), level=1)
        doc.add_paragraph(content)
        word_filename = filename.replace('.txt', '') + ".docx"
        doc.save(word_filename)

        # Create download link
        with open(word_filename, "rb") as file:
            b64 = base64.b64encode(file.read()).decode()
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="{word_filename}">📥 Download {word_filename}</a>'
        st.markdown(href, unsafe_allow_html=True)

    except Exception as e:
        st.error(f"Error creating file: {str(e)}")


# ---------- Existing Files Section ----------
st.subheader("📂 Existing Files in VM")
try:
    files = os.listdir(".")
    st.write(files)
except Exception as e:
    st.error(f"Error listing files: {str(e)}")


